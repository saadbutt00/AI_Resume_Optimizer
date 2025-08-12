# streamlit_resume_optimizer.py
# AI Resume Optimiser and Generator (Streamlit)
# Uses: pdfplumber, spaCy, python-docx, openai

import streamlit as st
import pdfplumber
import spacy
import re
import os
import io
import difflib
from docx import Document
import openai

# -------------------- CONFIG --------------------
# Make sure to set OPENAI_API_KEY environment variable or enter below
OPENAI_API_KEY = os.environ.get("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", "")
if OPENAI_API_KEY:
    openai.api_key = OPENAI_API_KEY

# spaCy model - recommended: en_core_web_md (has vectors)
try:
    nlp = spacy.load("en_core_web_md")
except Exception:
    # fallback to small model; quality for semantic similarity will be lower
    nlp = spacy.load("en_core_web_sm")

# -------------------- UTILITIES --------------------

def extract_text_from_pdf(file_bytes):
    """Extract text from uploaded PDF file bytes using pdfplumber."""
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n\n".join(text_chunks)


def simple_section_split(text):
    """Try to split a resume into common sections using headings heuristics."""
    headings = [r"experience", r"work experience", r"professional experience", r"education",
                r"skills", r"projects", r"certifications", r"summary", r"achievements", r"contact"]
    # Build regex to find headings
    pattern = r"(^|\n)\s*(?:" + "|".join(headings) + r")[:\n].*"
    parts = re.split(pattern, text, flags=re.IGNORECASE)
    # If split didn't work, fallback to whole text
    if len(parts) <= 1:
        return {"full_text": text}
    # Build a rough map: this is intentionally simple and may be improved
    sections = {"full_text": text}
    return sections


def extract_keywords_spacy(text, top_k=30):
    """Extract candidate keywords from text using spaCy noun chunks and entities."""
    doc = nlp(text)
    candidates = []
    for chunk in doc.noun_chunks:
        token_text = chunk.text.strip()
        if len(token_text) > 2:
            candidates.append(token_text.lower())
    for ent in doc.ents:
        if len(ent.text) > 2:
            candidates.append(ent.text.lower())
    # frequency score
    freq = {}
    for c in candidates:
        freq[c] = freq.get(c, 0) + 1
    # sort and return top_k
    items = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [k for k, v in items[:top_k]]


def compute_keyword_overlap(resume_text, jd_text):
    resume_kw = set(extract_keywords_spacy(resume_text, top_k=200))
    jd_kw = set(extract_keywords_spacy(jd_text, top_k=200))
    matched = resume_kw & jd_kw
    missing = jd_kw - resume_kw
    return {
        "resume_keywords_count": len(resume_kw),
        "jd_keywords_count": len(jd_kw),
        "matched_count": len(matched),
        "matched": sorted(list(matched))[:50],
        "missing": sorted(list(missing))[:50]
    }


def build_prompt_for_optimization(resume_text, jd_text, instructions=None):
    base = (
        "You are an expert resume writer and ATS optimisation assistant.\n"
        "Given the candidate resume and the target job description below, produce an optimised, ATS-friendly resume.\n"
        "Requirements:\n"
        "- Keep content truthful; do not invent qualifications.\n"
        "- Emphasise, rephrase, and reorder existing experiences to match the job description keywords.\n"
        "- Insert missing keywords only if they can be naturally applied (do not invent new roles).\n"
        "- Output the resume in plain text using clear headings: Contact, Summary, Skills, Experience, Education, Projects (if any).\n"
        "- Provide a brief change log after the resume: list added keywords and reworded sections.\n\n"
    )
    if instructions:
        base += "Extra instructions: " + instructions + "\n\n"
    base += "---BEGIN RESUME---\n" + resume_text + "\n---END RESUME---\n\n"
    base += "---BEGIN JOB DESCRIPTION---\n" + jd_text + "\n---END JOB DESCRIPTION---\n\n"
    base += "Produce only the optimised resume and then the change log.\n"
    return base


def call_openai_chat(prompt, model="gpt-4", max_tokens=1200, temperature=0.2):
    if not openai.api_key:
        raise RuntimeError("OpenAI API key not set. Set OPENAI_API_KEY environment variable or supply a key.")
    client = openai.OpenAI(api_key=openai.api_key)

    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens,
        temperature=temperature,
    )
    return resp.choices[0].message.content.strip()


def generate_docx_from_text(text):
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        elif re.match(r"^(Contact|Summary|Skills|Experience|Education|Projects|Certifications):", line):
            # heading
            doc.add_heading(line.strip(), level=2)
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def generate_change_summary(old_text, new_text):
    old_lines = old_text.splitlines()
    new_lines = new_text.splitlines()
    diff = difflib.unified_diff(old_lines, new_lines, lineterm="", n=0)
    return "\n".join(diff)

# -------------------- STREAMLIT UI --------------------

st.set_page_config(page_title="AI Resume Optimiser", layout="wide")
st.title("AI Resume Optimiser and Generator")
st.write("Upload a candidate resume (PDF) and paste a job description. The app will parse, analyse keywords, and call an LLM to produce an optimised resume and a change log.")

with st.sidebar:
    st.header("Settings")
    model_choice = st.selectbox("LLM model (requires API key)", ["gpt-4", "gpt-4o", "gpt-3.5-turbo"], index=0)
    openai_key_input = st.text_input("OpenAI API Key (or set env var OPENAI_API_KEY)", type="password")
    if openai_key_input:
        openai.api_key = openai_key_input

col1, col2 = st.columns([1, 2])

with col1:
    uploaded_file = st.file_uploader("Upload resume (PDF)", type=["pdf"], accept_multiple_files=False)
    jd_text = st.text_area("Paste job description here", height=300)
    instructions = st.text_area("Optional: Extra instructions for optimisation (tone, seniority etc)", height=120)
    optimize_btn = st.button("Optimise Resume")

with col2:
    st.subheader("Parsed Resume Text")
    parsed_text_area = st.empty()

resume_text = ""
if uploaded_file is not None:
    file_bytes = uploaded_file.read()
    try:
        resume_text = extract_text_from_pdf(file_bytes)
    except Exception as e:
        st.error(f"Failed to read PDF: {e}")
    parsed_text_area.code(resume_text[:10000] + ("..." if len(resume_text) > 10000 else ""))

if optimize_btn:
    if not resume_text:
        st.error("Please upload a resume PDF first.")
    elif not jd_text or jd_text.strip() == "":
        st.error("Please paste the target job description.")
    else:
        with st.spinner("Analysing and optimising..."):
            try:
                # Keyword overlap
                overlap = compute_keyword_overlap(resume_text, jd_text)

                st.subheader("Keyword Analysis")
                st.write(f"Resume keywords found: {overlap['resume_keywords_count']}")
                st.write(f"Job description keywords found: {overlap['jd_keywords_count']}")
                st.write(f"Matched keywords: {overlap['matched_count']}")
                st.write("Top matched keywords:")
                st.write(overlap['matched'])
                st.write("Top missing keywords from JD:")
                st.write(overlap['missing'])

                # Build prompt and call LLM
                prompt = build_prompt_for_optimization(resume_text, jd_text, instructions=instructions)
                llm_out = call_openai_chat(prompt, model=model_choice)

                # Split returned content: assume optimised resume then change log
                st.subheader("Optimised Resume (LLM Output)")
                st.code(llm_out[:10000] + ("..." if len(llm_out) > 10000 else ""))

                # Create downloadable DOCX
                docx_bytes = generate_docx_from_text(llm_out)
                st.download_button("Download Optimised Resume (DOCX)", data=docx_bytes.getvalue(), file_name="optimised_resume.docx")

                # Change summary using difflib (compare plain texts)
                st.subheader("Change Summary (diff)")
                change_summary = generate_change_summary(resume_text, llm_out)
                st.text_area("Diff (unified) between original resume text and optimised resume", value=change_summary, height=300)

            except Exception as e:
                st.error(f"Error during optimisation: {e}")

# Footer notes
st.markdown("---")
st.write("Notes: \n- This is a starter implementation. Improve parsing heuristics, add DOC/DOCX support, and tune the LLM prompt for best results.\n- Install spaCy model: `python -m spacy download en_core_web_md`.\n- Dependencies: pdfplumber, spacy, python-docx, openai, streamlit.")

# streamlit run "f:/app/app.py"